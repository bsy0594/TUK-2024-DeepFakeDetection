import streamlit as st
import time
import os
import tempfile
from docx import Document
from docx.shared import Inches

from sidebar import sidebar

st.write("# Create analytics reports")

sidebar()

st.session_state.documentation = True
high_prob_frames = st.session_state.high_prob_frames
probabilities = st.session_state.prediction
frame_indices = st.session_state.frame_index

high_prob_frames_sorted = sorted(high_prob_frames, key=lambda x: x[2], reverse=True)
top_3_frames = high_prob_frames_sorted[:3]
note = ["의심" if frame[2] > 0.5 else "정상" for frame in top_3_frames]

# 문서를 제작할 때 사용할 결과가 있는 경우
if "documentation" in st.session_state and st.session_state.documentation:
    # 로그인이 되어 있지 않은 경우
    if st.session_state["authentication_status"] == None:
        st.error("If you create a documentation while not signed in, it will not contain your information.")

    TEMPLATE_PATH = "doc_template.py"
    doc = Document(TEMPLATE_PATH)

    # 템플릿 내 텍스트 변경
    detection_results = {
        "보고서 발행일": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()),
        "분석 수행일": st.session_state.analysis_date,
        "파일명": st.session_state.uploaded_file.name,
        "파일 크기": f"{st.session_state.uploaded_file.size / 1048576:.2f} MB",
        "총 프레임 수": len(st.session_state.frame_index),
        "탐지 모델": st.session_state.model_name,
        "프레임 번호 1": top_3_frames[0][0],
        "탐지 확률 1": f"{top_3_frames[0][2]*100:.2f}%",
        "비고 1": note[0],
        "프레임 번호 2": top_3_frames[1][0],
        "탐지 확률 2": f"{top_3_frames[1][2]*100:.2f}%",
        "비고 2": note[1],
        "프레임 번호 3": top_3_frames[2][0],
        "탐지 확률 3": f"{top_3_frames[2][2]*100:.2f}%",
        "비고 3": note[2],
        "프레임 비율": f"{len(high_prob_frames) / len(frame_indices) * 100:.2f}%",
        "최고 확률": max(probabilities),
        "평균 확률": sum(probabilities) / len(probabilities),
    }
    img_bytes = st.session_state.img_bytes
    for para in doc.paragraphs:
        if "{{그래프}}" in para.text:
            para.text = ""
            run = para.add_run()
            run.add_picture(img_bytes, width=docx.shared.Inches(6))
        for key, value in detection_results.items():
            para.text = para.text.replace(f"{{{key}}}", str(value))

    # 수정된 문서 저장
    doc.save("deepfake_report.docx")

    # 사용자 다운로드
    with open("deepfake_report.docx", "rb") as f:
        st.download_button(
            label="Download report",
            data=f,
            file_name="deepfake_report.docx",
            mime="application/octet-stream"
        )
    
    
# 문서를 제작할 때 사용할 결과가 없는 경우
else:
    st.error("No previous detection results")
    st.error("You can view previous detection results on My Page")
    st.error("After a few moments, you will be redirected to My Page")
    time.sleep(3)
    st.switch_page("pages/My_Page.py")