import os
from urllib.parse import urljoin
import streamlit as st
import requests
import json
import plotly.express as px
import pandas as pd
from datetime import datetime
from PIL import Image
from io import BytesIO
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main_result(placeholder, uploaded_file, model_name):
    placeholder.empty()

    # 서버로 파일 및 옵션 전송
    FASTAPI_URL = "http://127.0.0.1:8000"
    detection_post_endpoint = "/video/"
    api_url = urljoin(FASTAPI_URL, detection_post_endpoint)

    files = {"file": uploaded_file.getvalue()}
    data = {"model": model_name}
    response = requests.post(api_url, files=files, data=data, timeout=60)

    if response.status_code == 200:
        # 서버 응답 저장
        data = response.json()
        
        frame_index = []
        original_image = []
        gradcam_image = []
        prediction = []

        for image in data["images"]:
            frame_index.append(image["frame_index"])
            original_image.append(image["original_image"])
            gradcam_image.append(image["gradcam_image"])
            prediction.append(image["prediction"])

        # 확률이 0.5 이상인 프레임이 있는 경우 딥페이크로 판단
        high_prob_frames = []
        max_prob_frame = None
        max_probability = 0

        for index, prob in enumerate(prediction):
            if prob > 0.5:
                high_prob_frames.append((original_image[index], gradcam_image[index],prob))
                if prob > max_probability:
                    max_probability = prob
                    max_prob_frame = original_image[index]

        st.session_state.high_prob_frames = high_prob_frames
        st.session_state.prediction = prediction
        st.session_state.frame_index = frame_index

        # 확률이 가장 높은 프레임을 메인으로 출력
        if len(high_prob_frames) > 0:
            st.markdown("# ⚠️ Deepfake is detected ⚠️")
            frame_url = urljoin(FASTAPI_URL, max_prob_frame)
            response = requests.get(frame_url)

            if response.status_code == 200:
                image = Image.open(BytesIO(response.content))
                st.image(image, use_container_width=True)
            else:
                st.error("Image not found")
        else:
            st.markdown("# No Deepfake Detected 🎉")

        # 결과 자세히 보러가기 버튼
        if "clicked" not in st.session_state:
            st.session_state.clicked = False

        def click_button():
            st.session_state.clicked = True

        st.button("View results in detail", on_click=click_button)

    else:
        st.error(f"Error: {response.status_code}")

# 문서 제작 기능
def generate_document():
    high_prob_frames = st.session_state.high_prob_frames
    probabilities = st.session_state.prediction
    frame_indices = st.session_state.frame_index

    high_prob_frames_sorted = sorted(high_prob_frames, key=lambda x: x[2], reverse=True)
    top_3_frames = high_prob_frames_sorted[:3]
    note = ["의심" if frame[2] > 0.5 else "정상" for frame in top_3_frames]


    TEMPLATE_PATH = "doc_template.docx"
    doc = Document(TEMPLATE_PATH)

    # 문서 작성 시 필요한 데이터
    detection_results = {
        "{{보고서 발행일}}": f"{time.strftime('%Y-%m-%d', time.localtime())}",
        "{{분석 수행일}}": f"{st.session_state.analysis_date}",
        "{{파일명}}": f"{st.session_state.uploaded_file.name}",
        "{{파일 크기}}": f"{st.session_state.uploaded_file.size / 1048576:.2f} MB",
        "{{총 프레임 수}}": f"{len(st.session_state.frame_index)}",
        "{{탐지 모델}}": f"{st.session_state.model_name}",
        "{{프레임 번호 1}}": f"{os.path.basename(top_3_frames[0][0])}",
        "{{탐지 확률 1}}": f"{top_3_frames[0][2]*100:.2f}%",
        "{{비고 1}}": note[0],
        "{{프레임 번호 2}}": f"{os.path.basename(top_3_frames[1][0])}",
        "{{탐지 확률 2}}": f"{top_3_frames[1][2]*100:.2f}%",
        "{{비고 2}}": note[1],
        "{{프레임 번호 3}}": f"{os.path.basename(top_3_frames[2][0])}",
        "{{탐지 확률 3}}": f"{top_3_frames[2][2]*100:.2f}%",
        "{{비고 3}}": note[2],
        "{{프레임 비율}}": f"{len(high_prob_frames) / len(frame_indices) * 100:.2f}% ({len(high_prob_frames)}개 / 총 {len(frame_indices)}개)",
        "{{최고 확률}}": f"{max(probabilities) * 100:.2f}%",
        "{{평균 확률}}": f"{sum(probabilities) / len(probabilities) * 100 :.2f}%",
    }
    img_bytes = st.session_state.img_bytes
    
    # XML 네임스페이스 정의 (Word의 'drawing' 요소를 찾기 위해 필요)
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    img_bytes = st.session_state.img_bytes

    # 표에서 텍스트 치환 및 그래프 삽입
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 1. "< 영상 프레임별 이상 탐지 확률 분석 >"이 포함된 셀 확인
                if "< 영상 프레임별 이상 탐지 확률 분석 >" in cell.text and img_bytes:
                    paragraph = cell.paragraphs[0]

                    # 2. 셀에 이미지가 이미 있는지 확인
                    has_image = any(run._element.find(".//w:drawing", nsmap) is not None for run in paragraph.runs)

                    if not has_image:
                        # 3. 이미지가 없으면 새로운 문단을 만들고 기존 문단 위에 삽입
                        new_paragraph = cell.add_paragraph()  # 새 문단 추가

                        # 🔹 문단 여백 제거 (공백 없애기)
                        # new_paragraph.paragraph_format.space_before = Pt(0)  # 문단 위 여백 제거
                        # new_paragraph.paragraph_format.space_after = Pt(0)   # 문단 아래 여백 제거
                        # new_paragraph.paragraph_format.line_spacing = Pt(1)  # 줄 간격 최소화

                        # 🔹 문단 중앙 정렬 설정 (이미지도 중앙 정렬됨)
                        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 

                        run = new_paragraph.add_run()
                        run.add_picture(img_bytes, width=Inches(5))  # 이미지 삽입

                        # 4. 새 문단을 기존 문단보다 위로 이동
                        cell._element.insert(0, new_paragraph._element)


                # 5. 텍스트 치환
                for key, value in detection_results.items():
                    if key in cell.text:
                        paragraph = cell.paragraphs[0]
                        paragraph.clear()  # 기존 내용 삭제
                        run = paragraph.add_run(str(value))  # 새 텍스트 추가
                        run.font.name = "Malgun Gothic"  # 기본 폰트 설정
                        run.font.size = Pt(9.5)  # 기본 글자 크기
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 중앙 정렬

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output

# 자세한 결과 출력
def detail_result(placeholder):
    placeholder.empty()

    high_prob_frames = st.session_state.high_prob_frames
    probabilities = st.session_state.prediction
    frame_indices = st.session_state.frame_index
    uploaded_file = st.session_state.uploaded_file
    model_name = st.session_state.model_name

    st.markdown("# ⚠️ Deepfake is detected ⚠️")

    # 영상, 모델, 시간 정보 출력
    col1, col2 = st.columns([2, 2])

    with col1:
        st.markdown(f"**Name:** {uploaded_file.name}")  
        st.markdown(f"**Size:** {uploaded_file.size / 1048576:.2f} MB")

    with col2:
        analysis_date = datetime.now().strftime('%Y-%m-%d')
        st.session_state.analysis_date = analysis_date
        st.markdown(f"**Model:** {model_name}")
        st.markdown(f"**Date:** {analysis_date}")

    # 확률이 높은 프레임들 출력
    st.markdown(
        """
        ######
        ### High Probability Frames
        """
    )
    FASTAPI_URL = "http://127.0.0.1:8000"

    frame_index = st.slider("Select Frame", 0, len(high_prob_frames) - 1, 0)
    original_image, gradcam_image, prob = high_prob_frames[frame_index]
    original_image_url = urljoin(FASTAPI_URL, original_image)
    gradcam_image_url = urljoin(FASTAPI_URL, gradcam_image)

    original_image_response = requests.get(original_image_url)
    gradcam_image_response = requests.get(gradcam_image_url)

    # gradcam 방식으로 프레임들 출력
    gradcam_toggle = st.checkbox("Show Grad-CAM")

    if gradcam_toggle:
        if gradcam_image_response.status_code == 200:
            frame_name = os.path.basename(gradcam_image_url)
            image = Image.open(BytesIO(gradcam_image_response.content))
            st.image(image, caption=f"'{frame_name}' is suspected to be deepfake with {prob*100:.2f}%")
        else:
            st.error("Gradcam Image not found")
        # frame_name = os.path.basename(gradcam_image_url)
        # st.image(gradcam_image_url, caption=f"'{frame_name}' is suspected to be deepfake with {prob*100:.2f}%")
    else:
        if original_image_response.status_code == 200:
            frame_name = os.path.basename(original_image_url)
            image = Image.open(BytesIO(original_image_response.content))
            st.image(image, caption=f"'{frame_name}' is suspected to be deepfake with {prob*100:.2f}%")
        else:
            st.error("Original Image not found")

    # 프레임 별 확률에 대한 그래프
    st.markdown("### ")
    st.markdown("### Deepfake Probability per Frame")

    df = pd.DataFrame({
        "Frame": frame_indices,
        "Probability": probabilities
    })

    fig = px.line(df, x="Frame", y="Probability")
    fig.update_xaxes(rangeslider_visible=True)
    st.plotly_chart(fig, use_container_width=True)
    
    # 문서 제작 시 필요한 그래프 만들기
    doc_fig = px.line(df, x="Frame", y="Probability")

    # 🔹 상하좌우 여백 최소화
    doc_fig.update_layout(
        margin=dict(l=50, r=10, t=10, b=40),  # 좌우상하 마진 최소화
        height=300,  # 그래프 높이 조정
        # xaxis_title=None,  # X축 제목 제거 (선택 사항)
        # yaxis_title=None,  # Y축 제목 제거 (선택 사항)
        title=None  # 제목 제거 (선택 사항)
    )

    img_bytes = BytesIO()
    doc_fig.write_image(img_bytes, format="png")
    img_bytes.seek(0)
    st.session_state.img_bytes = img_bytes

    # 문서 제작하러 가기
    doc_bytes = generate_document()
    st.download_button(
        label="Download Report",
        data=doc_bytes,
        file_name="deepfake_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )